import streamlit as st
import pandas as pd
from docx import Document
import subprocess
import os
import tempfile

st.set_page_config(page_title="Dispute Document Generator", layout="centered")

LOGO_URL = "https://raw.githubusercontent.com/vijayp-alt/dispute-doc-generator/main/image.webp"

# ── Glass / Frosted Theme ────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

* { font-family: 'Inter', sans-serif; }

[data-testid="stAppViewContainer"] {
    background: linear-gradient(135deg, #dbeafe 0%, #bfdbfe 40%, #e0f2fe 100%);
    min-height: 100vh;
}

[data-testid="stAppViewContainer"]::before {
    content: '';
    position: fixed;
    top: -50%;
    left: -50%;
    width: 200%;
    height: 200%;
    background:
        radial-gradient(ellipse at 20% 20%, rgba(99,102,241,0.08) 0%, transparent 50%),
        radial-gradient(ellipse at 80% 80%, rgba(59,130,246,0.08) 0%, transparent 50%),
        radial-gradient(ellipse at 50% 50%, rgba(147,197,253,0.1) 0%, transparent 60%);
    pointer-events: none;
    z-index: 0;
}

[data-testid="stHeader"] { background: transparent !important; }
#MainMenu, footer { visibility: hidden; }
[data-testid="stMain"] { position: relative; z-index: 1; }

/* ── Glass Card ── */
.glass-card {
    background: rgba(255,255,255,0.55);
    backdrop-filter: blur(20px);
    -webkit-backdrop-filter: blur(20px);
    border: 1px solid rgba(255,255,255,0.8);
    border-radius: 20px;
    padding: 28px 32px;
    margin-bottom: 20px;
}

/* ── Header ── */
.glass-header {
    background: rgba(255,255,255,0.6);
    backdrop-filter: blur(30px);
    -webkit-backdrop-filter: blur(30px);
    border: 1px solid rgba(255,255,255,0.85);
    border-radius: 20px;
    padding: 22px 28px;
    margin-bottom: 24px;
    display: flex;
    align-items: center;
    gap: 18px;
}

.header-title {
    color: #1e3a5f;
    font-size: 1.75rem;
    font-weight: 700;
    margin: 0;
    letter-spacing: -0.5px;
}

.header-subtitle {
    color: #4a6fa5;
    font-size: 0.82rem;
    margin: 3px 0 0 0;
}

/* ── Progress Steps ── */
.steps-wrap {
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 0;
    margin-bottom: 24px;
}
.step {
    display: flex;
    align-items: center;
    gap: 8px;
}
.step-circle {
    width: 32px;
    height: 32px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 0.78rem;
    font-weight: 600;
    border: 1.5px solid rgba(99,102,241,0.25);
    color: #93a8d4;
    background: rgba(255,255,255,0.4);
}
.step-circle.active {
    background: linear-gradient(135deg, #6366f1, #a855f7);
    border-color: transparent;
    color: white;
}
.step-circle.done {
    background: rgba(99,102,241,0.15);
    border-color: #6366f1;
    color: #6366f1;
}
.step-label {
    font-size: 0.78rem;
    color: #93a8d4;
    font-weight: 500;
}
.step-label.active { color: #6366f1; }
.step-label.done   { color: #818cf8; }
.step-line {
    width: 48px;
    height: 1.5px;
    background: rgba(99,102,241,0.15);
    margin: 0 4px;
}
.step-line.done { background: #6366f1; }

/* ── Card title ── */
.card-title {
    color: #1e3a5f;
    font-size: 1rem;
    font-weight: 600;
    margin-bottom: 16px;
}

/* ── Section label ── */
.section-label {
    color: #4a6fa5;
    font-size: 0.72rem;
    font-weight: 600;
    letter-spacing: 1.1px;
    text-transform: uppercase;
    margin-bottom: 8px;
}

/* ── Summary pill ── */
.summary-pill {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: rgba(99,102,241,0.1);
    border: 1px solid rgba(99,102,241,0.3);
    border-radius: 999px;
    padding: 5px 14px;
    font-size: 0.8rem;
    color: #6366f1;
    margin-bottom: 14px;
}

/* ── File uploader ── */
[data-testid="stFileUploader"] {
    background: rgba(255,255,255,0.4) !important;
    border: 1.5px dashed rgba(99,102,241,0.35) !important;
    border-radius: 14px !important;
}
[data-testid="stFileUploader"]:hover {
    border-color: rgba(99,102,241,0.6) !important;
}

/* ── Generate button ── */
div[data-testid="stButton"] > button {
    background: linear-gradient(135deg, #6366f1 0%, #a855f7 100%);
    color: white !important;
    border: none !important;
    border-radius: 14px;
    padding: 14px 32px;
    font-size: 1rem;
    font-weight: 600;
    width: 100%;
    cursor: pointer;
    position: relative;
    overflow: hidden;
    animation: pulse-glow 2.5s infinite;
}
@keyframes pulse-glow {
    0%   { box-shadow: 0 0 0 0 rgba(99,102,241,0.5); }
    50%  { box-shadow: 0 0 20px 6px rgba(168,85,247,0.35); }
    100% { box-shadow: 0 0 0 0 rgba(99,102,241,0.5); }
}
div[data-testid="stButton"] > button:hover {
    transform: translateY(-2px);
    box-shadow: 0 8px 30px rgba(99,102,241,0.5) !important;
}

/* ── Download button ── */
div[data-testid="stDownloadButton"] > button {
    background: rgba(16,185,129,0.2) !important;
    border: 1px solid rgba(16,185,129,0.5) !important;
    color: #6ee7b7 !important;
    border-radius: 14px;
    padding: 12px 28px;
    font-size: 0.95rem;
    font-weight: 600;
    width: 100%;
}
div[data-testid="stDownloadButton"] > button:hover {
    background: rgba(16,185,129,0.35) !important;
}

/* ── Dataframe ── */
[data-testid="stDataFrame"] {
    border-radius: 12px;
    overflow: hidden;
    border: 1px solid rgba(255,255,255,0.08) !important;
}

/* ── Alerts ── */
[data-testid="stAlert"] {
    background: rgba(255,255,255,0.5) !important;
    border: 1px solid rgba(99,102,241,0.2) !important;
    border-radius: 12px !important;
    color: #1e3a5f !important;
}

p, label, span, div { color: #1e3a5f; }
</style>
""", unsafe_allow_html=True)

# ── Confetti JS ──────────────────────────────────────────────────────────────
CONFETTI_JS = """
<script src="https://cdn.jsdelivr.net/npm/canvas-confetti@1.6.0/dist/confetti.browser.min.js"></script>
<script>
(function(){
    var end = Date.now() + 2500;
    var colors = ['#6366f1','#a855f7','#ec4899','#06b6d4','#c4b5fd'];
    (function frame(){
        confetti({particleCount:3, angle:60, spread:55, origin:{x:0}, colors:colors});
        confetti({particleCount:3, angle:120, spread:55, origin:{x:1}, colors:colors});
        if(Date.now()<end) requestAnimationFrame(frame);
    }());
})();
</script>
"""

# ── Header ───────────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="glass-header">
    <img src="{LOGO_URL}" width="64" onerror="this.style.display='none'"
         style="border-radius:12px; border:1px solid rgba(255,255,255,0.15);"/>
    <div>
        <p class="header-title">Dispute Document Generator</p>
        <p class="header-subtitle">Upload your Excel data and Word template to generate a dispute PDF instantly.</p>
    </div>
</div>
""", unsafe_allow_html=True)

# ── Determine current step ────────────────────────────────────────────────────
# We'll track step via session state
if "pdf_ready" not in st.session_state:
    st.session_state.pdf_ready = False

# ── Progress Indicator ────────────────────────────────────────────────────────
def render_steps(step):
    def sc(n): return "done" if n < step else ("active" if n == step else "")
    def sl(n): return "done" if n < step else ("active" if n == step else "")
    def ln(n): return "done" if n >= step else ""
    st.markdown(f"""
    <div class="steps-wrap">
      <div class="step">
        <div class="step-circle {sc(1)}">{'✓' if step>1 else '1'}</div>
        <span class="step-label {sl(1)}">Upload</span>
      </div>
      <div class="step-line {ln(2)}"></div>
      <div class="step">
        <div class="step-circle {sc(2)}">{'✓' if step>2 else '2'}</div>
        <span class="step-label {sl(2)}">Preview</span>
      </div>
      <div class="step-line {ln(3)}"></div>
      <div class="step">
        <div class="step-circle {sc(3)}">{'✓' if step>3 else '3'}</div>
        <span class="step-label {sl(3)}">Generate</span>
      </div>
      <div class="step-line {ln(4)}"></div>
      <div class="step">
        <div class="step-circle {sc(4)}">{'✓' if step>4 else '4'}</div>
        <span class="step-label {sl(4)}">Download</span>
      </div>
    </div>
    """, unsafe_allow_html=True)

# ── Upload Card ───────────────────────────────────────────────────────────────
render_steps(1)

st.markdown('<div class="glass-card">', unsafe_allow_html=True)
st.markdown('<div class="card-title">📂 Upload Files</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    st.markdown('<div class="section-label">Excel Data File</div>', unsafe_allow_html=True)
    excel_file = st.file_uploader("", type=["xls", "xlsx"], key="excel")
with col2:
    st.markdown('<div class="section-label">Word Template</div>', unsafe_allow_html=True)
    word_file = st.file_uploader("", type=["docx"], key="word")

st.markdown('</div>', unsafe_allow_html=True)

# ── Processing ────────────────────────────────────────────────────────────────
if excel_file and word_file:
    try:
        df = pd.read_excel(excel_file, sheet_name=0, engine='openpyxl')

        # ── Step 2: Preview ──
        render_steps(2)
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">📊 Data Preview</div>', unsafe_allow_html=True)

        if 'Field name' in df.columns and 'Value' in df.columns:
            field_count = len(df)
            st.markdown(f'<div class="summary-pill">📋 {field_count} fields detected — ready to populate</div>', unsafe_allow_html=True)

        st.dataframe(df, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

        if 'Field name' in df.columns and 'Value' in df.columns:
            field_map = dict(zip(df['Field name'], df['Value']))

            # ── Step 3: Generate ──
            render_steps(3)
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.markdown('<div class="card-title">⚙️ Generate Document</div>', unsafe_allow_html=True)

            if st.button("🚀 Generate PDF"):
                with st.spinner("✨ Building your dispute document..."):
                    with tempfile.TemporaryDirectory() as tmpdir:
                        temp_doc_path = os.path.join(tmpdir, "updated.docx")

                        doc = Document(word_file)

                        for paragraph in doc.paragraphs:
                            for key, value in field_map.items():
                                placeholder = f"<{key}>"
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, str(value))

                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for key, value in field_map.items():
                                        placeholder = f"<{key}>"
                                        if placeholder in cell.text:
                                            cell.text = cell.text.replace(placeholder, str(value))

                        doc.save(temp_doc_path)

                        result = subprocess.run(
                            ["libreoffice", "--headless", "--convert-to", "pdf",
                             "--outdir", tmpdir, temp_doc_path],
                            capture_output=True, text=True
                        )
                        if result.returncode != 0:
                            raise Exception(f"LibreOffice conversion failed: {result.stderr}")

                        temp_pdf_path = os.path.join(tmpdir, "updated.pdf")

                        with open(temp_pdf_path, "rb") as f:
                            pdf_bytes = f.read()

                st.session_state.pdf_ready = True
                st.session_state.pdf_bytes = pdf_bytes

            st.markdown('</div>', unsafe_allow_html=True)

            # ── Step 4: Download ──
            if st.session_state.get("pdf_ready"):
                render_steps(4)
                st.markdown('<div class="glass-card">', unsafe_allow_html=True)
                st.markdown('<div class="card-title">🎉 Your PDF is Ready!</div>', unsafe_allow_html=True)
                st.markdown(CONFETTI_JS, unsafe_allow_html=True)
                st.download_button(
                    "📥 Download Dispute PDF",
                    data=st.session_state.pdf_bytes,
                    file_name="Dispute_Document.pdf",
                    mime="application/pdf"
                )
                st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.error("⚠️ Excel file must contain 'Field name' and 'Value' columns.")

    except Exception as e:
        st.error(f"❌ Error: {e}")

else:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.info("👆 Please upload both an Excel file and a Word template above to get started.")
    st.markdown('</div>', unsafe_allow_html=True)
